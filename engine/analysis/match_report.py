"""
SpaceAI FC - Match Summary Generator
======================================
Combines all analysis modules into one structured match report.
This is the central piece that ties the engine together.

Phase 2 additions:
    - Formation detection via FormationDetector (replaces simple method)
    - Player role classification
    - Press resistance analysis
    - Tactical pattern detection
"""

import matplotlib.pyplot as plt
import matplotlib.gridspec as gridspec
from matplotlib.patches import FancyBboxPatch
from mplsoccer import Pitch
import numpy as np
from datetime import datetime


class MatchReport:
    """
    Generates a complete match analysis report by combining:
        - Player positions and formations
        - Pass network analysis
        - Space control analysis
        - Formation detection (Phase 2)
        - Player roles (Phase 2)
        - Press resistance (Phase 2)
        - Tactical patterns (Phase 2)
    
    Outputs:
        - Structured text report
        - Visual dashboard combining all analyses
        - Word document report
    """
    
    def __init__(self):
        self.team_a = {'name': 'Team A', 'color': '#e74c3c', 'players': []}
        self.team_b = {'name': 'Team B', 'color': '#3498db', 'players': []}
        self.ball_pos = None
        self.pass_network = None
        self.space_control = None
        self.match_info = {}
        
        # Phase 2 modules
        self.formation_detector_a = None
        self.formation_detector_b = None
        self.role_classifier_a = None
        self.role_classifier_b = None
        self.press_resistance = None
        self.pattern_detector = None
    
    # ── Setup ───────────────────────────────────────────────────
    
    def set_match_info(self, home_team, away_team, score_home=0, score_away=0,
                       minute=0, competition="Friendly", date=None):
        """Set general match information."""
        self.match_info = {
            'home_team': home_team,
            'away_team': away_team,
            'score_home': score_home,
            'score_away': score_away,
            'minute': minute,
            'competition': competition,
            'date': date or datetime.now().strftime("%Y-%m-%d"),
        }
    
    def set_team_a(self, name, color, players):
        """Set team A data."""
        self.team_a = {'name': name, 'color': color, 'players': players}
    
    def set_team_b(self, name, color, players):
        """Set team B data."""
        self.team_b = {'name': name, 'color': color, 'players': players}
    
    def set_ball(self, x, y):
        """Set ball position."""
        self.ball_pos = (x, y)
    
    def set_pass_network(self, pass_network):
        """Attach a computed PassNetwork object."""
        self.pass_network = pass_network
    
    def set_space_control(self, space_control):
        """Attach a computed SpaceControl object."""
        self.space_control = space_control
    
    # Phase 2 setters
    
    def set_formation_detector(self, detector_a, detector_b=None):
        """Attach FormationDetector objects for both teams."""
        self.formation_detector_a = detector_a
        self.formation_detector_b = detector_b
    
    def set_role_classifier(self, classifier_a, classifier_b=None):
        """Attach RoleClassifier objects for both teams."""
        self.role_classifier_a = classifier_a
        self.role_classifier_b = classifier_b
    
    def set_press_resistance(self, press_resistance):
        """Attach a PressResistance object."""
        self.press_resistance = press_resistance
    
    def set_pattern_detector(self, pattern_detector):
        """Attach a PatternDetector object."""
        self.pattern_detector = pattern_detector
    
    # ── Formation Detection ─────────────────────────────────────
    
    def _detect_formation(self, players, detector=None):
        """
        Detect formation using FormationDetector if available,
        otherwise fall back to simple gap-based method.
        """
        if detector is not None:
            result = detector.detect()
            return result['formation']
        return self._detect_formation_simple(players)
    
    def _detect_formation_simple(self, players):
        """
        Simple formation detection by grouping players into lines
        based on x-coordinate clustering.
        
        Detects which side the team is on and reads defense-to-attack.
        This is a basic version — Phase 2 FormationDetector is preferred.
        """
        outfield = [p for p in players if p.get('position') != 'GK']
        
        if not outfield:
            return "Unknown"
        
        sorted_players = sorted(outfield, key=lambda p: p['x'])
        x_positions = [p['x'] for p in sorted_players]
        
        avg_x = sum(x_positions) / len(x_positions)
        
        gaps = []
        for i in range(1, len(x_positions)):
            gaps.append((x_positions[i] - x_positions[i-1], i))
        
        gaps.sort(reverse=True)
        
        for num_splits in [3, 2]:
            if len(gaps) >= num_splits:
                split_positions = sorted([idx for _, idx in gaps[:num_splits]])
                
                lines = []
                prev = 0
                for sp in split_positions:
                    lines.append(sp - prev)
                    prev = sp
                lines.append(len(x_positions) - prev)
                
                if all(1 <= l <= 5 for l in lines) and sum(lines) == 10:
                    if avg_x > 60:
                        lines.reverse()
                    return "-".join(str(l) for l in lines)
        
        if len(gaps) >= 2:
            split_positions = sorted([idx for _, idx in gaps[:2]])
            lines = []
            prev = 0
            for sp in split_positions:
                lines.append(sp - prev)
                prev = sp
            lines.append(len(x_positions) - prev)
            if avg_x > 60:
                lines.reverse()
            return "-".join(str(l) for l in lines)
        
        return "Unknown"
    
    # ── Generate full report data ───────────────────────────────
    
    def generate_report(self):
        """
        Generate a complete structured report combining all analyses.
        
        Returns a dict with all findings.
        """
        report = {
            'match_info': self.match_info,
            'team_a': {
                'name': self.team_a['name'],
                'formation': self._detect_formation(
                    self.team_a['players'], self.formation_detector_a),
                'player_count': len(self.team_a['players']),
            },
            'team_b': {
                'name': self.team_b['name'],
                'formation': self._detect_formation(
                    self.team_b['players'], self.formation_detector_b),
                'player_count': len(self.team_b['players']),
            },
            'pass_analysis': None,
            'space_analysis': None,
            'roles_a': None,
            'roles_b': None,
            'press_resistance': None,
            'patterns_a': None,
            'patterns_b': None,
            'insights': [],
            'recommendations': [],
        }
        
        # Formation confidence (Phase 2)
        if self.formation_detector_a:
            r = self.formation_detector_a.detect()
            report['team_a']['formation_confidence'] = r['confidence']
        if self.formation_detector_b:
            r = self.formation_detector_b.detect()
            report['team_b']['formation_confidence'] = r['confidence']
        
        # Pass network analysis
        if self.pass_network:
            summary = self.pass_network.get_summary()
            report['pass_analysis'] = summary
            
            kp = summary['key_distributor']
            report['insights'].append(
                f"{kp['name']} (#{kp['number']}) is the key distributor "
                f"with betweenness centrality of {kp['betweenness']}"
            )
            
            mi = summary['most_involved']
            report['insights'].append(
                f"{mi['name']} (#{mi['number']}) is the most involved player "
                f"with {mi['passes_made']} passes made and {mi['passes_received']} received"
            )
            
            if summary['top_connections']:
                top = summary['top_connections'][0]
                report['insights'].append(
                    f"Strongest connection: {top['from']} -> {top['to']} "
                    f"({top['passes']} passes)"
                )
            
            if summary['weak_links']:
                for wl in summary['weak_links']:
                    report['insights'].append(
                        f"Weak link: {wl['name']} (#{wl['number']}) with only "
                        f"{wl['total_involvement']} total pass involvements"
                    )
                    report['recommendations'].append(
                        f"Increase involvement of {wl['name']} in build-up play"
                    )
        
        # Space control analysis
        if self.space_control:
            control_grid, stats = self.space_control.compute_influence_control()
            midfield = self.space_control.get_midfield_control(control_grid)
            
            report['space_analysis'] = {
                'overall': stats,
                'midfield': midfield,
            }
            
            ta = self.team_a['name']
            tb = self.team_b['name']
            
            if stats['team_a_control'] > 55:
                report['insights'].append(f"{ta} is dominating territorial control ({stats['team_a_control']}%)")
            elif stats['team_b_control'] > 55:
                report['insights'].append(f"{tb} is dominating territorial control ({stats['team_b_control']}%)")
            else:
                report['insights'].append(f"Space control is balanced ({ta} {stats['team_a_control']}% vs {tb} {stats['team_b_control']}%)")
            
            if midfield['team_a'] > 55:
                report['insights'].append(f"{ta} controls midfield ({midfield['team_a']}%)")
            elif midfield['team_b'] > 55:
                report['insights'].append(f"{tb} controls midfield ({midfield['team_b']}%)")
                report['recommendations'].append(f"{ta} should add an extra midfielder or drop deeper to contest midfield")
            
            zones = stats['zones']
            if zones['attacking_third']['team_a'] < 20:
                report['recommendations'].append(
                    f"{ta} has very low presence in the attacking third — "
                    f"push forward or increase width"
                )
            if zones['middle_third']['team_b'] > 60:
                report['recommendations'].append(
                    f"{tb} is dominating the middle third — "
                    f"{ta} should press higher or switch play faster"
                )
        
        # ── Phase 2: Player Roles ───────────────────────────────
        if self.role_classifier_a:
            roles = self.role_classifier_a.classify_all()
            report['roles_a'] = roles
            # Add notable role insights
            for r in roles:
                if r['role'] in ('False Nine', 'Sweeper Keeper', 'Inverted Winger',
                                 'Inverted Fullback', 'Shadow Striker'):
                    report['insights'].append(
                        f"{r['name']} operating as {r['role']} ({r['confidence']:.0%} conf.)"
                    )
        
        if self.role_classifier_b:
            roles = self.role_classifier_b.classify_all()
            report['roles_b'] = roles
        
        # ── Phase 2: Press Resistance ───────────────────────────
        if self.press_resistance:
            pr_result = self.press_resistance.analyze()
            report['press_resistance'] = pr_result
            
            score = pr_result['press_resistance_score']
            team = self.press_resistance.team_name
            
            if score >= 70:
                report['insights'].append(
                    f"{team} shows strong press resistance (score: {score:.0f}/100)"
                )
            elif score < 45:
                report['insights'].append(
                    f"{team} is vulnerable under pressure (score: {score:.0f}/100)"
                )
                report['recommendations'].append(
                    f"{team} should vary build-up routes to avoid press traps"
                )
            
            if pr_result['vulnerable_zones']:
                worst = pr_result['vulnerable_zones'][0]
                report['recommendations'].append(
                    f"Avoid building through {worst['zone']} "
                    f"(only {worst['success_rate']:.0%} success rate)"
                )
        
        # ── Phase 2: Tactical Patterns ──────────────────────────
        if self.pattern_detector:
            patterns_a = self.pattern_detector.detect_all(team="a")
            patterns_b = self.pattern_detector.detect_all(team="b")
            report['patterns_a'] = patterns_a
            report['patterns_b'] = patterns_b
            
            ta = self.team_a['name']
            tb = self.team_b['name']
            
            for p in patterns_a:
                if p['detected']:
                    report['insights'].append(
                        f"{ta}: {p['pattern']} detected ({p['confidence']:.0%} conf.)"
                    )
            for p in patterns_b:
                if p['detected']:
                    report['insights'].append(
                        f"{tb}: {p['pattern']} detected ({p['confidence']:.0%} conf.)"
                    )
        
        if not report['recommendations']:
            report['recommendations'].append("Maintain current tactical approach — balanced performance detected")
        
        return report
    
    # ── Text output ─────────────────────────────────────────────
    
    def print_report(self):
        """Print a complete formatted match report."""
        report = self.generate_report()
        mi = report['match_info']
        
        print("\n")
        print("+" + "=" * 60 + "+")
        print("|" + " " * 10 + "SPACEAI FC - MATCH ANALYSIS REPORT" + " " * 16 + "|")
        print("+" + "=" * 60 + "+")
        
        print(f"\n  {mi.get('competition', 'Match')}")
        print(f"  {mi.get('home_team', 'Home')} {mi.get('score_home', 0)} - {mi.get('score_away', 0)} {mi.get('away_team', 'Away')}")
        print(f"  Date: {mi.get('date', 'N/A')}    Minute: {mi.get('minute', 0)}'")
        
        print(f"\n  --- FORMATIONS ---")
        conf_a = report['team_a'].get('formation_confidence', '')
        conf_b = report['team_b'].get('formation_confidence', '')
        conf_a_str = f" ({conf_a:.0%} conf.)" if conf_a else ""
        conf_b_str = f" ({conf_b:.0%} conf.)" if conf_b else ""
        print(f"  {report['team_a']['name']}: {report['team_a']['formation']}{conf_a_str}")
        print(f"  {report['team_b']['name']}: {report['team_b']['formation']}{conf_b_str}")
        
        if report['pass_analysis']:
            pa = report['pass_analysis']
            print(f"\n  --- PASS ANALYSIS ---")
            print(f"  Total Passes: {pa['total_passes']}")
            print(f"  Key Distributor: {pa['key_distributor']['name']} (#{pa['key_distributor']['number']})")
            print(f"  Most Involved: {pa['most_involved']['name']} (#{pa['most_involved']['number']})")
            print(f"  Top Connection: {pa['top_connections'][0]['from']} -> {pa['top_connections'][0]['to']} ({pa['top_connections'][0]['passes']} passes)")
        
        if report['space_analysis']:
            sa = report['space_analysis']
            overall = sa['overall']
            mid = sa['midfield']
            print(f"\n  --- SPACE CONTROL ---")
            print(f"  {report['team_a']['name']}: {overall['team_a_control']}%")
            print(f"  {report['team_b']['name']}: {overall['team_b_control']}%")
            print(f"  Midfield: {report['team_a']['name']} {mid['team_a']}% | {report['team_b']['name']} {mid['team_b']}%")
        
        # Phase 2 sections
        if report['roles_a']:
            print(f"\n  --- PLAYER ROLES ({report['team_a']['name']}) ---")
            for r in report['roles_a']:
                print(f"  #{r['number']:2d} {r['name']:15s} -> {r['role']}")
        
        if report['press_resistance']:
            pr = report['press_resistance']
            print(f"\n  --- PRESS RESISTANCE ---")
            print(f"  Score: {pr['press_resistance_score']:.0f}/100")
            print(f"  Success Under Pressure: {pr['pass_success_under_pressure']:.0%}")
            print(f"  Escape Rate: {pr['escape_rate']:.0%}")
        
        if report['patterns_a']:
            detected = [p for p in report['patterns_a'] if p['detected']]
            if detected:
                print(f"\n  --- TACTICAL PATTERNS ({report['team_a']['name']}) ---")
                for p in detected:
                    print(f"  [+] {p['pattern']} ({p['confidence']:.0%})")
        
        if report['patterns_b']:
            detected = [p for p in report['patterns_b'] if p['detected']]
            if detected:
                print(f"\n  --- TACTICAL PATTERNS ({report['team_b']['name']}) ---")
                for p in detected:
                    print(f"  [+] {p['pattern']} ({p['confidence']:.0%})")
        
        print(f"\n  --- KEY INSIGHTS ---")
        for i, insight in enumerate(report['insights'], 1):
            print(f"  {i}. {insight}")
        
        print(f"\n  --- TACTICAL RECOMMENDATIONS ---")
        for i, rec in enumerate(report['recommendations'], 1):
            print(f"  {i}. {rec}")
        
        print("\n+" + "=" * 60 + "+")
        print("|" + " " * 18 + "End of Report" + " " * 29 + "|")
        print("+" + "=" * 60 + "+\n")
    
    # ── Visual dashboard ────────────────────────────────────────
    
    def draw_dashboard(self, figsize=(20, 14)):
        """
        Generate a visual dashboard combining:
            - Top left: pitch with players
            - Top right: pass network
            - Bottom left: space control
            - Bottom right: stats and insights text
        """
        fig = plt.figure(figsize=figsize, facecolor='#1a1a2e')
        gs = gridspec.GridSpec(2, 2, figure=fig, hspace=0.3, wspace=0.2)
        
        report = self.generate_report()
        mi = report['match_info']
        
        title = f"SpaceAI FC Match Analysis: {mi.get('home_team', 'Home')} vs {mi.get('away_team', 'Away')}"
        fig.suptitle(title, color='white', fontsize=18, fontweight='bold', y=0.98)
        
        ax1 = fig.add_subplot(gs[0, 0])
        self._draw_pitch_panel(ax1, report)
        
        ax2 = fig.add_subplot(gs[0, 1])
        self._draw_pass_panel(ax2)
        
        ax3 = fig.add_subplot(gs[1, 0])
        self._draw_space_panel(ax3)
        
        ax4 = fig.add_subplot(gs[1, 1])
        self._draw_stats_panel(ax4, report)
        
        return fig
    
    def _draw_pitch_panel(self, ax, report=None):
        """Draw pitch with both teams and formation labels."""
        pitch = Pitch(pitch_type='statsbomb', pitch_color='#1a1a2e',
                      line_color='#e0e0e0', linewidth=1, goal_type='box')
        pitch.draw(ax=ax)
        
        ax.set_title("Formation & Positions", color='white', fontsize=12,
                      fontweight='bold', pad=10)
        
        for p in self.team_a['players']:
            ax.scatter(p['x'], p['y'], c=self.team_a['color'], s=200,
                       edgecolors='white', linewidths=1.5, zorder=6)
            ax.annotate(str(p['number']), xy=(p['x'], p['y']),
                       ha='center', va='center', fontsize=7,
                       fontweight='bold', color='white', zorder=7)
        
        for p in self.team_b['players']:
            ax.scatter(p['x'], p['y'], c=self.team_b['color'], s=200,
                       edgecolors='white', linewidths=1.5, zorder=6)
            ax.annotate(str(p['number']), xy=(p['x'], p['y']),
                       ha='center', va='center', fontsize=7,
                       fontweight='bold', color='white', zorder=7)
        
        if self.ball_pos:
            ax.scatter(self.ball_pos[0], self.ball_pos[1], c='#f1c40f', s=80,
                       edgecolors='white', linewidths=1.5, zorder=8)
        
        # Use Phase 2 detector if available
        if report:
            form_a = report['team_a']['formation']
            form_b = report['team_b']['formation']
        else:
            form_a = self._detect_formation(self.team_a['players'], self.formation_detector_a)
            form_b = self._detect_formation(self.team_b['players'], self.formation_detector_b)
        
        ax.text(25, -3, f"{self.team_a['name']}: {form_a}", ha='center',
                color=self.team_a['color'], fontsize=9, fontweight='bold')
        ax.text(95, -3, f"{self.team_b['name']}: {form_b}", ha='center',
                color=self.team_b['color'], fontsize=9, fontweight='bold')
    
    def _draw_pass_panel(self, ax):
        """Draw pass network panel."""
        pitch = Pitch(pitch_type='statsbomb', pitch_color='#1a1a2e',
                      line_color='#e0e0e0', linewidth=1, goal_type='box')
        pitch.draw(ax=ax)
        
        ax.set_title("Pass Network", color='white', fontsize=12,
                      fontweight='bold', pad=10)
        
        if not self.pass_network:
            ax.text(60, 40, "No pass data", ha='center', va='center',
                    color='#888888', fontsize=14)
            return
        
        pn = self.pass_network
        counts = pn.get_pass_counts()
        max_weight = max((d['weight'] for _, _, d in pn.graph.edges(data=True)), default=1)
        max_inv = max((c['total_involvement'] for c in counts.values()), default=1)
        
        for passer, receiver, data in pn.graph.edges(data=True):
            if data['weight'] >= 2 and passer in pn.players and receiver in pn.players:
                x1 = pn.players[passer]['x']
                y1 = pn.players[passer]['y']
                x2 = pn.players[receiver]['x']
                y2 = pn.players[receiver]['y']
                
                w = data['weight']
                lw = 0.5 + (w / max_weight) * 4
                alpha = 0.3 + (w / max_weight) * 0.5
                
                ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="-|>", color=self.team_a['color'],
                                    lw=lw, alpha=alpha,
                                    connectionstyle="arc3,rad=0.15",
                                    mutation_scale=12), zorder=3)
        
        for pid, player in pn.players.items():
            inv = counts[pid]['total_involvement']
            size = 150 + (inv / max_inv) * 350
            ax.scatter(player['x'], player['y'], c=self.team_a['color'], s=size,
                       edgecolors='white', linewidths=1.5, zorder=6)
            ax.annotate(str(player['number']), xy=(player['x'], player['y']),
                       ha='center', va='center', fontsize=7,
                       fontweight='bold', color='white', zorder=7)
    
    def _draw_space_panel(self, ax):
        """Draw space control panel."""
        pitch = Pitch(pitch_type='statsbomb', pitch_color='#1a1a2e',
                      line_color='#e0e0e0', linewidth=1, goal_type='box')
        pitch.draw(ax=ax)
        
        ax.set_title("Space Control", color='white', fontsize=12,
                      fontweight='bold', pad=10)
        
        if not self.space_control:
            ax.text(60, 40, "No space data", ha='center', va='center',
                    color='#888888', fontsize=14)
            return
        
        from matplotlib.colors import LinearSegmentedColormap
        
        control_grid, stats = self.space_control.compute_influence_control()
        
        cmap = LinearSegmentedColormap.from_list('ctrl',
            [self.team_a['color'], '#1a1a2e', self.team_b['color']], N=256)
        
        extent = [0, 120, 0, 80]
        ax.imshow(control_grid, extent=extent, origin='lower',
                  cmap=cmap, alpha=0.5, aspect='auto', zorder=2,
                  vmin=-1, vmax=1)
        
        for p in self.team_a['players']:
            ax.scatter(p['x'], p['y'], c=self.team_a['color'], s=150,
                       edgecolors='white', linewidths=1.5, zorder=6)
        for p in self.team_b['players']:
            ax.scatter(p['x'], p['y'], c=self.team_b['color'], s=150,
                       edgecolors='white', linewidths=1.5, zorder=6)
        
        ax.text(60, -3,
                f"{self.team_a['name']}: {stats['team_a_control']}% | "
                f"{self.team_b['name']}: {stats['team_b_control']}%",
                ha='center', color='white', fontsize=9, fontweight='bold')
    
    def _draw_stats_panel(self, ax, report=None):
        """Draw stats and insights text panel (with Phase 2 data)."""
        if report is None:
            report = self.generate_report()
        
        ax.set_facecolor('#1a1a2e')
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 10)
        ax.axis('off')
        
        ax.set_title("Analysis Summary", color='white', fontsize=12,
                      fontweight='bold', pad=10)
        
        mi = report['match_info']
        
        y = 9.5
        line_height = 0.38
        
        score_text = f"{mi.get('home_team', 'Home')} {mi.get('score_home', 0)} - {mi.get('score_away', 0)} {mi.get('away_team', 'Away')}"
        ax.text(5, y, score_text, ha='center', va='top', fontsize=13,
                fontweight='bold', color='white')
        y -= line_height * 1.5
        
        ax.text(5, y, f"{mi.get('competition', '')}  |  {mi.get('minute', 0)}'",
                ha='center', va='top', fontsize=9, color='#aaaaaa')
        y -= line_height * 1.5
        
        # Formations
        ax.text(0.5, y, "FORMATIONS", va='top', fontsize=9,
                fontweight='bold', color='#f1c40f')
        y -= line_height
        ax.text(0.5, y, f"{report['team_a']['name']}: {report['team_a']['formation']}",
                va='top', fontsize=8, color='white')
        y -= line_height
        ax.text(0.5, y, f"{report['team_b']['name']}: {report['team_b']['formation']}",
                va='top', fontsize=8, color='white')
        y -= line_height * 1.2
        
        # Press Resistance (Phase 2)
        if report['press_resistance']:
            pr = report['press_resistance']
            ax.text(0.5, y, "PRESS RESISTANCE", va='top', fontsize=9,
                    fontweight='bold', color='#f1c40f')
            y -= line_height
            score_val = pr['press_resistance_score']
            score_color = '#2ecc71' if score_val >= 65 else '#f39c12' if score_val >= 45 else '#e74c3c'
            ax.text(0.5, y, f"Score: {score_val:.0f}/100  |  "
                    f"Pressure success: {pr['pass_success_under_pressure']:.0%}",
                    va='top', fontsize=8, color=score_color)
            y -= line_height * 1.2
        
        # Key Insights
        ax.text(0.5, y, "KEY INSIGHTS", va='top', fontsize=9,
                fontweight='bold', color='#f1c40f')
        y -= line_height
        
        for insight in report['insights'][:5]:
            display = insight[:55] + "..." if len(insight) > 55 else insight
            ax.text(0.5, y, display, va='top', fontsize=7, color='white')
            y -= line_height
        
        y -= line_height * 0.3
        
        # Recommendations
        ax.text(0.5, y, "RECOMMENDATIONS", va='top', fontsize=9,
                fontweight='bold', color='#2ecc71')
        y -= line_height
        
        for rec in report['recommendations'][:3]:
            display = rec[:55] + "..." if len(rec) > 55 else rec
            ax.text(0.5, y, "• " + display, va='top', fontsize=7, color='white')
            y -= line_height
    
    # ── Save ────────────────────────────────────────────────────
    
    def save_dashboard(self, fig, filename="match_report.png"):
        """Save the dashboard to outputs folder."""
        fig.savefig(f"outputs/{filename}", dpi=150, bbox_inches='tight',
                    facecolor=fig.get_facecolor())
        print(f"Saved: outputs/{filename}")
    
    # ── Word document export ────────────────────────────────────
    
    def export_document(self, filename="match_report.docx"):
        """Export the match report as a Word document with Phase 2 sections."""
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import os
        
        doc = DocxDocument()
        
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        report = self.generate_report()
        mi = report['match_info']
        
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("SpaceAI FC - Match Analysis Report")
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0x1A, 0x35, 0x50)
        
        # Match Info
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(
            f"{mi.get('home_team', 'Home')} {mi.get('score_home', 0)} - "
            f"{mi.get('score_away', 0)} {mi.get('away_team', 'Away')}"
        )
        run.bold = True
        run.font.size = Pt(18)
        
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = info.add_run(
            f"{mi.get('competition', '')}  |  Minute: {mi.get('minute', 0)}'  |  {mi.get('date', '')}"
        )
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        
        doc.add_paragraph("")
        
        # Formations
        h = doc.add_heading("Formations", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x1A, 0x35, 0x50)
        
        conf_a = report['team_a'].get('formation_confidence', '')
        conf_b = report['team_b'].get('formation_confidence', '')
        conf_a_str = f" (confidence: {conf_a:.0%})" if conf_a else ""
        conf_b_str = f" (confidence: {conf_b:.0%})" if conf_b else ""
        doc.add_paragraph(f"{report['team_a']['name']}: {report['team_a']['formation']}{conf_a_str}")
        doc.add_paragraph(f"{report['team_b']['name']}: {report['team_b']['formation']}{conf_b_str}")
        
        # Pass Analysis
        if report['pass_analysis']:
            pa = report['pass_analysis']
            h = doc.add_heading("Pass Analysis", level=1)
            h.runs[0].font.color.rgb = RGBColor(0x1A, 0x35, 0x50)
            
            doc.add_paragraph(f"Total Passes: {pa['total_passes']}")
            doc.add_paragraph(
                f"Key Distributor: {pa['key_distributor']['name']} "
                f"(#{pa['key_distributor']['number']})"
            )
            doc.add_paragraph(
                f"Most Involved: {pa['most_involved']['name']} "
                f"(#{pa['most_involved']['number']}) - "
                f"{pa['most_involved']['passes_made']} made, "
                f"{pa['most_involved']['passes_received']} received"
            )
            
            h2 = doc.add_heading("Top Connections", level=2)
            h2.runs[0].font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
            for conn in pa['top_connections']:
                doc.add_paragraph(
                    f"{conn['from']} \u2192 {conn['to']}: {conn['passes']} passes",
                    style='List Bullet'
                )
            
            if pa['weak_links']:
                h2 = doc.add_heading("Weak Links", level=2)
                h2.runs[0].font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
                for wl in pa['weak_links']:
                    doc.add_paragraph(
                        f"{wl['name']} (#{wl['number']}): {wl['total_involvement']} total involvements",
                        style='List Bullet'
                    )
        
        # Space Control
        if report['space_analysis']:
            sa = report['space_analysis']
            overall = sa['overall']
            mid = sa['midfield']
            
            h = doc.add_heading("Space Control", level=1)
            h.runs[0].font.color.rgb = RGBColor(0x1A, 0x35, 0x50)
            
            doc.add_paragraph(
                f"Overall: {report['team_a']['name']} {overall['team_a_control']}% | "
                f"{report['team_b']['name']} {overall['team_b_control']}%"
            )
            doc.add_paragraph(
                f"Midfield Control: {report['team_a']['name']} {mid['team_a']}% | "
                f"{report['team_b']['name']} {mid['team_b']}%"
            )
            
            zones = overall['zones']
            h2 = doc.add_heading("Zone Breakdown", level=2)
            h2.runs[0].font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
            doc.add_paragraph(
                f"Defensive Third: {report['team_a']['name']} {zones['defensive_third']['team_a']}% | "
                f"{report['team_b']['name']} {zones['defensive_third']['team_b']}%",
                style='List Bullet'
            )
            doc.add_paragraph(
                f"Middle Third: {report['team_a']['name']} {zones['middle_third']['team_a']}% | "
                f"{report['team_b']['name']} {zones['middle_third']['team_b']}%",
                style='List Bullet'
            )
            doc.add_paragraph(
                f"Attacking Third: {report['team_a']['name']} {zones['attacking_third']['team_a']}% | "
                f"{report['team_b']['name']} {zones['attacking_third']['team_b']}%",
                style='List Bullet'
            )
        
        # ── Phase 2: Player Roles ───────────────────────────────
        if report['roles_a'] or report['roles_b']:
            h = doc.add_heading("Player Roles", level=1)
            h.runs[0].font.color.rgb = RGBColor(0x1A, 0x35, 0x50)
            
            if report['roles_a']:
                h2 = doc.add_heading(report['team_a']['name'], level=2)
                h2.runs[0].font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
                for r in report['roles_a']:
                    doc.add_paragraph(
                        f"#{r['number']} {r['name']}: {r['role']} ({r['confidence']:.0%})",
                        style='List Bullet'
                    )
            
            if report['roles_b']:
                h2 = doc.add_heading(report['team_b']['name'], level=2)
                h2.runs[0].font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
                for r in report['roles_b']:
                    doc.add_paragraph(
                        f"#{r['number']} {r['name']}: {r['role']} ({r['confidence']:.0%})",
                        style='List Bullet'
                    )
        
        # ── Phase 2: Press Resistance ───────────────────────────
        if report['press_resistance']:
            pr = report['press_resistance']
            h = doc.add_heading("Press Resistance", level=1)
            h.runs[0].font.color.rgb = RGBColor(0x1A, 0x35, 0x50)
            
            score = pr['press_resistance_score']
            rating = "Excellent" if score >= 75 else "Good" if score >= 60 else "Average" if score >= 45 else "Poor"
            
            doc.add_paragraph(f"Press Resistance Score: {score:.0f}/100 ({rating})")
            doc.add_paragraph(f"Passes Under Pressure: {pr['passes_under_pressure']}/{pr['total_passes']}")
            doc.add_paragraph(f"Success Under Pressure: {pr['pass_success_under_pressure']:.0%}")
            doc.add_paragraph(f"Escape Rate: {pr['escape_rate']:.0%}")
            
            if pr['vulnerable_zones']:
                h2 = doc.add_heading("Vulnerable Zones", level=2)
                h2.runs[0].font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
                for vz in pr['vulnerable_zones']:
                    doc.add_paragraph(
                        f"{vz['zone']}: {vz['success_rate']:.0%} success ({vz['total_passes']} passes)",
                        style='List Bullet'
                    )
        
        # ── Phase 2: Tactical Patterns ──────────────────────────
        if report['patterns_a'] or report['patterns_b']:
            h = doc.add_heading("Tactical Patterns", level=1)
            h.runs[0].font.color.rgb = RGBColor(0x1A, 0x35, 0x50)
            
            for team_key, team_lbl in [('patterns_a', report['team_a']['name']),
                                        ('patterns_b', report['team_b']['name'])]:
                patterns = report.get(team_key) or []
                detected = [p for p in patterns if p['detected']]
                if detected:
                    h2 = doc.add_heading(team_lbl, level=2)
                    h2.runs[0].font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
                    for p in detected:
                        doc.add_paragraph(
                            f"{p['pattern']} ({p['confidence']:.0%}): {p['description']}",
                            style='List Bullet'
                        )
        
        # Visualizations
        h = doc.add_heading("Visualizations", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x1A, 0x35, 0x50)
        
        image_files = [
            ("outputs/01_pitch.png", "Formation & Player Positions"),
            ("outputs/02_pass_network.png", "Pass Network"),
            ("outputs/03_pass_sequence.png", "Build-Up Sequence"),
            ("outputs/04_space_voronoi.png", "Space Control (Voronoi)"),
            ("outputs/05_space_influence.png", "Space Control (Influence)"),
            ("outputs/06_formation_barca.png", "Formation Detection (Barcelona)"),
            ("outputs/06_formation_madrid.png", "Formation Detection (Real Madrid)"),
            ("outputs/07_roles_barca.png", "Player Roles (Barcelona)"),
            ("outputs/07_roles_madrid.png", "Player Roles (Real Madrid)"),
            ("outputs/08_press_resistance.png", "Press Resistance"),
            ("outputs/09_patterns_barca.png", "Tactical Patterns (Barcelona)"),
            ("outputs/09_patterns_madrid.png", "Tactical Patterns (Real Madrid)"),
            ("outputs/10_match_dashboard.png", "Full Match Dashboard"),
        ]
        
        for img_path, caption in image_files:
            if os.path.exists(img_path):
                h2 = doc.add_heading(caption, level=2)
                h2.runs[0].font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
                doc.add_picture(img_path, width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph("")
        
        # Key Insights
        h = doc.add_heading("Key Insights", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x1A, 0x35, 0x50)
        
        for i, insight in enumerate(report['insights'], 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. ")
            run.bold = True
            p.add_run(insight)
        
        # Recommendations
        h = doc.add_heading("Tactical Recommendations", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x1A, 0x35, 0x50)
        
        for rec in report['recommendations']:
            doc.add_paragraph(rec, style='List Bullet')
        
        # Footer
        doc.add_paragraph("")
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Generated by SpaceAI FC - Tactical Analysis Engine v2")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        run.italic = True
        
        filepath = f"outputs/{filename}"
        doc.save(filepath)
        print(f"Saved: {filepath}")


# ═══════════════════════════════════════════════════════════════
# Quick test - Full El Clásico analysis
# ═══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    
    import sys
    sys.path.insert(0, '.')
    from engine.analysis.pass_network import PassNetwork
    from engine.analysis.space_control import SpaceControl
    from engine.analysis.formation_detection import FormationDetector
    from engine.analysis.role_classifier import RoleClassifier
    from engine.analysis.press_resistance import PressResistance
    from engine.analysis.pattern_detection import PatternDetector
    import random
    
    barca_players = [
        {'name': 'ter Stegen',  'number': 1,  'x': 5,  'y': 40, 'position': 'GK'},
        {'name': 'Koundé',      'number': 23, 'x': 30, 'y': 70, 'position': 'RB'},
        {'name': 'Araújo',      'number': 4,  'x': 25, 'y': 52, 'position': 'CB'},
        {'name': 'Cubarsí',     'number': 2,  'x': 25, 'y': 28, 'position': 'CB'},
        {'name': 'Baldé',       'number': 3,  'x': 30, 'y': 10, 'position': 'LB'},
        {'name': 'Pedri',       'number': 8,  'x': 45, 'y': 48, 'position': 'CM'},
        {'name': 'De Jong',     'number': 21, 'x': 45, 'y': 32, 'position': 'CM'},
        {'name': 'Lamine',      'number': 19, 'x': 65, 'y': 68, 'position': 'RW'},
        {'name': 'Gavi',        'number': 6,  'x': 60, 'y': 40, 'position': 'CAM'},
        {'name': 'Raphinha',    'number': 11, 'x': 65, 'y': 12, 'position': 'LW'},
        {'name': 'Lewandowski', 'number': 9,  'x': 80, 'y': 40, 'position': 'ST'},
    ]
    
    madrid_players = [
        {'name': 'Courtois',    'number': 1,  'x': 115, 'y': 40, 'position': 'GK'},
        {'name': 'Carvajal',    'number': 2,  'x': 90,  'y': 70, 'position': 'RB'},
        {'name': 'Rüdiger',     'number': 22, 'x': 93,  'y': 52, 'position': 'CB'},
        {'name': 'Militão',     'number': 3,  'x': 93,  'y': 28, 'position': 'CB'},
        {'name': 'Mendy',       'number': 23, 'x': 90,  'y': 10, 'position': 'LB'},
        {'name': 'Tchouaméni',  'number': 14, 'x': 73,  'y': 40, 'position': 'CDM'},
        {'name': 'Valverde',    'number': 15, 'x': 70,  'y': 58, 'position': 'CM'},
        {'name': 'Bellingham',  'number': 5,  'x': 70,  'y': 22, 'position': 'CM'},
        {'name': 'Rodrygo',     'number': 11, 'x': 55,  'y': 65, 'position': 'RW'},
        {'name': 'Mbappé',      'number': 7,  'x': 50,  'y': 40, 'position': 'ST'},
        {'name': 'Vinícius',    'number': 20, 'x': 55,  'y': 15, 'position': 'LW'},
    ]
    
    pn = PassNetwork()
    pn.add_players(barca_players)
    
    passes = [
        (1, 4), (1, 2), (1, 4),
        (4, 8), (4, 21), (4, 8), (4, 23), (4, 2),
        (2, 21), (2, 3), (2, 8), (2, 21), (2, 4),
        (23, 19), (23, 8), (23, 19), (23, 4),
        (3, 11), (3, 21), (3, 11), (3, 2),
        (8, 6), (8, 19), (8, 23), (8, 21), (8, 9),
        (8, 6), (8, 19), (8, 4), (8, 9), (8, 11),
        (8, 6), (8, 21), (8, 23),
        (21, 8), (21, 2), (21, 3), (21, 6), (21, 11),
        (21, 8), (21, 4), (21, 6),
        (6, 9), (6, 19), (6, 8), (6, 11), (6, 9),
        (6, 8), (6, 19), (6, 21),
        (19, 9), (19, 23), (19, 6), (19, 8),
        (19, 9), (19, 6),
        (11, 9), (11, 3), (11, 6), (11, 21),
        (11, 9), (11, 6),
        (9, 6), (9, 19), (9, 8), (9, 11),
        (6, 8), (19, 8), (9, 6), (11, 21),
    ]
    pn.add_passes(passes)
    
    sc = SpaceControl()
    sc.set_teams(barca_players, madrid_players)
    sc.set_ball(60, 40)
    
    # Phase 2 modules
    fd_b = FormationDetector()
    fd_b.set_team(barca_players, "FC Barcelona", "#a50044")
    fd_m = FormationDetector()
    fd_m.set_team(madrid_players, "Real Madrid", "#ffffff")
    
    rc_b = RoleClassifier()
    rc_b.set_team(barca_players, "FC Barcelona", "#a50044")
    rc_m = RoleClassifier()
    rc_m.set_team(madrid_players, "Real Madrid", "#ffffff")
    
    random.seed(42)
    events = []
    outfield = [p for p in barca_players if p['position'] != 'GK']
    for _ in range(65):
        passer = random.choice(outfield)
        receiver = random.choice(outfield)
        px = max(0, min(120, passer['x'] + random.gauss(0, 3)))
        py = max(0, min(80, passer['y'] + random.gauss(0, 3)))
        opp_pos = np.array([[p['x'], p['y']] for p in madrid_players])
        dist = np.linalg.norm(opp_pos - np.array([px, py]), axis=1)
        nearby = np.sum(dist < 10)
        success = random.random() < (0.55 if nearby >= 2 else 0.90)
        events.append({'passer': passer['number'], 'receiver': receiver['number'],
                       'success': success, 'x': px, 'y': py,
                       'end_x': receiver['x'], 'end_y': receiver['y']})
    
    pr = PressResistance()
    pr.set_teams(barca_players, madrid_players,
                 "FC Barcelona", "#a50044", "Real Madrid")
    pr.add_pass_events(events)
    
    ptd = PatternDetector()
    ptd.set_teams(barca_players, madrid_players,
                  "FC Barcelona", "Real Madrid", "#a50044", "#ffffff")
    
    mr = MatchReport()
    mr.set_match_info(home_team="FC Barcelona", away_team="Real Madrid",
                       score_home=2, score_away=1, minute=72,
                       competition="La Liga", date="2026-03-22")
    mr.set_team_a("FC Barcelona", "#a50044", barca_players)
    mr.set_team_b("Real Madrid", "#ffffff", madrid_players)
    mr.set_ball(60, 40)
    mr.set_pass_network(pn)
    mr.set_space_control(sc)
    mr.set_formation_detector(fd_b, fd_m)
    mr.set_role_classifier(rc_b, rc_m)
    mr.set_press_resistance(pr)
    mr.set_pattern_detector(ptd)
    
    mr.print_report()
    
    fig = mr.draw_dashboard()
    mr.save_dashboard(fig, "el_clasico_report.png")
    mr.export_document("el_clasico_report.docx")
    
    plt.show()
    print("\nMatch report complete!")